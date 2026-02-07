"""Document parser for PDF and DOCX files."""

from pathlib import Path
from typing import Dict, Any

import PyPDF2
import docx

from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentParser:
    """Parse PDF and Word documents."""

    @staticmethod
    def parse_pdf(file_path: Path) -> Dict[str, Any]:
        """
        Parse PDF document.

        Args:
            file_path: Path to PDF file

        Returns:
            Parsed document data
        """
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"

                result = {
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "num_pages": len(reader.pages),
                    "text": text.strip(),
                    "metadata": reader.metadata if reader.metadata else {}
                }

                logger.info(f"Parsed PDF: {file_path.name} ({len(reader.pages)} pages)")
                return result

        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return {
                "file_name": file_path.name,
                "file_type": "pdf",
                "error": str(e)
            }

    @staticmethod
    def parse_docx(file_path: Path) -> Dict[str, Any]:
        """
        Parse Word document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Parsed document data
        """
        try:
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    tables_text.append(row_text)

            result = {
                "file_name": file_path.name,
                "file_type": "docx",
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "text": text,
                "tables": "\n".join(tables_text) if tables_text else None,
                "metadata": {
                    "core_properties": {
                        "author": doc.core_properties.author,
                        "created": str(doc.core_properties.created),
                        "modified": str(doc.core_properties.modified),
                        "title": doc.core_properties.title,
                    }
                }
            }

            logger.info(f"Parsed DOCX: {file_path.name} ({len(doc.paragraphs)} paragraphs)")
            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {
                "file_name": file_path.name,
                "file_type": "docx",
                "error": str(e)
            }

    @classmethod
    def parse(cls, file_path: Path) -> Dict[str, Any]:
        """
        Auto-detect and parse document.

        Args:
            file_path: Path to document file

        Returns:
            Parsed document data
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return cls.parse_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return cls.parse_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {suffix}")
            return {
                "file_name": file_path.name,
                "error": f"Unsupported file type: {suffix}"
            }
